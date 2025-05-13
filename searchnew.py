import streamlit as st
import requests
from bs4 import BeautifulSoup
import re
import time
import os
from docx import Document
from docx.shared import Inches

# 画像ダウンロード用関数
def download_image(url, folder="images"):
    if not os.path.exists(folder):
        os.makedirs(folder)
    filename = os.path.join(folder, url.split("/")[-1])
    response = requests.get(url, stream=True)
    if response.status_code == 200:
        with open(filename, "wb") as file:
            for chunk in response.iter_content(1024):
                file.write(chunk)
        return filename
    return None

# 検索してURLを取得する関数
def search_and_scrape(search_query):
    search_query = search_query.strip().replace(' ', '%20')
    result_links = []
    page_num = 1
    pattern = re.compile(r'/([1-9][0-9]{2,})[A-Za-z]\d{2}')
    
    while page_num <= 6:
        if page_num == 1:
            url = f'https://medu4.com/quizzes/result?q={search_query}&st=all'
        else:
            url = f'https://medu4.com/quizzes/result?page={page_num}&q={search_query}&st=all'

        response = requests.get(url)
        if response.status_code != 200:
            break

        soup = BeautifulSoup(response.text, 'html.parser')
        all_links = [link['href'] for link in soup.find_all('a', href=True)]
        page_results = [link for link in all_links if pattern.search(link)]

        if not page_results:
            break

        result_links.extend(page_results)
        page_num += 1
        time.sleep(0.5)

    full_urls = [f"https://medu4.com{link}" for link in result_links]
    return full_urls

# ページ内容を取得（画像あり／なし切替対応）
def get_page_text(url, get_images=True):
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')

    category = soup.find('span', class_='button-small-line')
    category_name = category.text.strip() if category else '分野名なし'

    problem = soup.find('div', class_='quiz-body mb-64')
    problem_text = problem.text.strip() if problem else '問題文なし'

    choices = []
    for choice in soup.find_all('div', class_='box-select'):
        choice_header = choice.find('span', {'class': 'choice-header'}).text.strip()
        choice_text = choice.find_all('span')[1].text.strip()
        choices.append(f"{choice_header} {choice_text}")

    h4_tags = soup.find_all('h4')
    answer_text = '解答なし'
    question_id = '問題番号なし'
    if len(h4_tags) >= 2:
        answer_text = h4_tags[0].text.strip()
        question_id_match = re.search(r'([0-9]{3}[A-Za-z][0-9]+)', h4_tags[1].text)
        if question_id_match:
            question_id = question_id_match.group(1)

    explanation = soup.find('div', class_='explanation')
    explanation_text = explanation.text.strip() if explanation else '解説なし'

    image_urls = []
    if get_images:
        image_divs = soup.find_all('div', class_='box-quiz-image mb-32')
        for div in image_divs:
            img_tag = div.find('img')
            if img_tag and img_tag.get('src'):
                img_url = img_tag['src']
                img_url_full = img_url.replace('thumb_', '')
                image_urls.append(img_url_full)

    return {
        "category": category_name,
        "problem": problem_text,
        "choices": choices,
        "answer": answer_text,
        "question_id": question_id,
        "explanation": explanation_text,
        "images": image_urls
    }

# Word出力（画像あり対応）
def create_word_doc(pages_data, search_query, include_images=True):
    doc = Document()
    doc.add_heading('検索結果', 0)
    doc.add_paragraph(f"取得問題数: {len(pages_data)}問")

    for idx, page_data in enumerate(pages_data, start=1):
        title = f"問題{idx} {page_data['question_id']}"
        doc.add_paragraph(title, style='Heading2')
        doc.add_paragraph(f"問題文: {page_data['problem']}")

        if include_images and page_data['images']:
            for img_url in page_data['images']:
                img_path = download_image(img_url)
                if img_path:
                    doc.add_paragraph()
                    doc.add_picture(img_path, width=Inches(2.5))
                else:
                    doc.add_paragraph(f"画像取得失敗: {img_url}")

        doc.add_paragraph("選択肢:")
        for choice in page_data['choices']:
            doc.add_paragraph(choice)
        doc.add_paragraph(f"{page_data['answer']}")
        doc.add_paragraph(f"解説: {page_data['explanation']}")
        doc.add_paragraph("-" * 50)

    filename = f"{search_query}_search_results.docx"
    doc.save(filename)
    return filename

# Streamlit UI
st.title("Medu4 検索ツールNew2")
search_query = st.text_input("検索ワードを入力してください")

col1, col2 = st.columns(2)

def run_search(get_images: bool):
    with st.spinner("検索中..."):
        result_pages = search_and_scrape(search_query)

    if result_pages:
        st.write(f"{len(result_pages)}件の問題が見つかりました。")
        progress_bar = st.progress(0)
        status_text = st.empty()

        pages_data = []
        for i, url in enumerate(result_pages):
            page_data = get_page_text(url, get_images=get_images)
            pages_data.append(page_data)

            progress = int((i + 1) / len(result_pages) * 100)
            progress_bar.progress(progress)
            status_text.text(f"{i + 1} / {len(result_pages)} 件取得中...")

        with st.spinner("Wordファイル作成中..."):
            filename = create_word_doc(pages_data, search_query, include_images=get_images)

        st.success("Wordファイルが完成しました！")
        with open(filename, "rb") as file:
            st.download_button("📄 Wordファイルをダウンロード", file, file_name=filename)
    else:
        st.error("検索結果がありませんでした。")

with col1:
    if st.button("🔍 検索（画像あり）"):
        run_search(get_images=True)

with col2:
    if st.button("⚡ 検索（画像なし）"):
        run_search(get_images=False)
